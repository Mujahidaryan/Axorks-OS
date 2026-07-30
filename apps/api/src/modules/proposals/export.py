"""
Axorks OS — Proposal PDF & DOCX Export

Generates formatted proposal documents from structured JSONB content.
"""

from __future__ import annotations

import io
from datetime import date, datetime
from decimal import Decimal
from typing import Any
from uuid import UUID

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from docx import Document
from docx.shared import Inches, Pt

from src.modules.proposals.models import Proposal


def _fmt_currency(value: Decimal | float | int | None, currency: str = "USD") -> str:
    if value is None:
        return "—"
    symbol = "$" if currency == "USD" else f"{currency} "
    return f"{symbol}{float(value):,.2f}"


def _safe_text(value: Any) -> str:
    if value is None:
        return ""
    return str(value).replace("\r\n", "\n")


def generate_proposal_pdf(proposal: Proposal) -> bytes:
    """Render proposal content as a PDF byte stream."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ProposalTitle",
        parent=styles["Heading1"],
        fontSize=20,
        spaceAfter=12,
        textColor=colors.HexColor("#1e1b4b"),
    )
    heading_style = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        fontSize=13,
        spaceBefore=14,
        spaceAfter=6,
        textColor=colors.HexColor("#312e81"),
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        spaceAfter=8,
    )
    meta_style = ParagraphStyle(
        "Meta",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.grey,
        alignment=TA_CENTER,
    )

    content = proposal.content or {}
    story: list[Any] = []

    story.append(Paragraph("AXORKS OS — PROPOSAL", meta_style))
    story.append(Paragraph(_safe_text(proposal.title), title_style))
    story.append(
        Paragraph(
            f"Type: {proposal.type.replace('_', ' ').title()} &nbsp;|&nbsp; "
            f"Date: {proposal.created_at.strftime('%B %d, %Y') if isinstance(proposal.created_at, datetime) else date.today()}",
            meta_style,
        )
    )
    if proposal.valid_until:
        story.append(Paragraph(f"Valid Until: {proposal.valid_until}", meta_style))
    story.append(Spacer(1, 0.25 * inch))

    sections = sorted(content.get("sections") or [], key=lambda s: s.get("order", 0))
    for section in sections:
        story.append(Paragraph(_safe_text(section.get("title", "Section")), heading_style))
        for paragraph in _safe_text(section.get("content", "")).split("\n"):
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), body_style))

    pricing = content.get("pricing") or {}
    items = pricing.get("items") or []
    if items:
        story.append(Paragraph("Investment Summary", heading_style))
        table_data = [["Description", "Qty", "Unit Price", "Amount"]]
        for item in items:
            table_data.append([
                _safe_text(item.get("description", "")),
                str(item.get("quantity", 1)),
                _fmt_currency(item.get("unit_price"), proposal.currency),
                _fmt_currency(item.get("amount"), proposal.currency),
            ])
        total = pricing.get("total") or proposal.total_value
        table_data.append(["", "", "Total", _fmt_currency(total, proposal.currency)])

        table = Table(table_data, colWidths=[3.2 * inch, 0.6 * inch, 1.0 * inch, 1.0 * inch])
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#ede9fe")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1e1b4b")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -2), 0.5, colors.lightgrey),
            ("LINEABOVE", (0, -1), (-1, -1), 1, colors.HexColor("#6366f1")),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("ALIGN", (1, 0), (-1, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        story.append(table)
        story.append(Spacer(1, 0.15 * inch))

    timeline = content.get("timeline") or {}
    milestones = timeline.get("milestones") or []
    if milestones:
        story.append(Paragraph("Project Timeline", heading_style))
        for milestone in milestones:
            line = f"<b>{_safe_text(milestone.get('title', 'Milestone'))}</b>"
            if milestone.get("duration"):
                line += f" — {_safe_text(milestone['duration'])}"
            if milestone.get("description"):
                line += f"<br/>{_safe_text(milestone['description'])}"
            story.append(Paragraph(line, body_style))

    payment_plan = content.get("payment_plan") or {}
    payment_milestones = payment_plan.get("milestones") or []
    if payment_milestones:
        story.append(Paragraph("Payment Plan", heading_style))
        for pm in payment_milestones:
            amount = _fmt_currency(pm.get("amount"), proposal.currency)
            pct = f" ({pm['percentage']}%)" if pm.get("percentage") else ""
            due = f" — Due: {pm['due_date']}" if pm.get("due_date") else ""
            story.append(
                Paragraph(f"• {_safe_text(pm.get('title', 'Payment'))}: {amount}{pct}{due}", body_style)
            )

    terms = content.get("terms_and_conditions")
    if terms:
        story.append(Paragraph("Terms & Conditions", heading_style))
        for paragraph in _safe_text(terms).split("\n"):
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), body_style))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def generate_proposal_docx(proposal: Proposal) -> bytes:
    """Render proposal content as a DOCX byte stream."""
    document = Document()
    content = proposal.content or {}

    title = document.add_heading(proposal.title, level=0)
    title.alignment = 1

    meta = document.add_paragraph(
        f"{proposal.type.replace('_', ' ').title()} • "
        f"{proposal.created_at.strftime('%B %d, %Y') if isinstance(proposal.created_at, datetime) else date.today()}"
    )
    meta.alignment = 1
    for run in meta.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = None

    sections = sorted(content.get("sections") or [], key=lambda s: s.get("order", 0))
    for section in sections:
        document.add_heading(_safe_text(section.get("title", "Section")), level=2)
        for paragraph in _safe_text(section.get("content", "")).split("\n"):
            if paragraph.strip():
                document.add_paragraph(paragraph.strip())

    pricing = content.get("pricing") or {}
    items = pricing.get("items") or []
    if items:
        document.add_heading("Investment Summary", level=2)
        table = document.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Description"
        hdr[1].text = "Qty"
        hdr[2].text = "Unit Price"
        hdr[3].text = "Amount"
        for item in items:
            row = table.add_row().cells
            row[0].text = _safe_text(item.get("description", ""))
            row[1].text = str(item.get("quantity", 1))
            row[2].text = _fmt_currency(item.get("unit_price"), proposal.currency)
            row[3].text = _fmt_currency(item.get("amount"), proposal.currency)
        total_row = table.add_row().cells
        total_row[0].text = ""
        total_row[1].text = ""
        total_row[2].text = "Total"
        total_row[3].text = _fmt_currency(pricing.get("total") or proposal.total_value, proposal.currency)

    timeline = content.get("timeline") or {}
    milestones = timeline.get("milestones") or []
    if milestones:
        document.add_heading("Project Timeline", level=2)
        for milestone in milestones:
            p = document.add_paragraph()
            p.add_run(_safe_text(milestone.get("title", "Milestone"))).bold = True
            if milestone.get("duration"):
                p.add_run(f" — {_safe_text(milestone['duration'])}")
            if milestone.get("description"):
                document.add_paragraph(_safe_text(milestone["description"]))

    payment_plan = content.get("payment_plan") or {}
    if payment_plan.get("milestones"):
        document.add_heading("Payment Plan", level=2)
        for pm in payment_plan["milestones"]:
            document.add_paragraph(
                f"{_safe_text(pm.get('title', 'Payment'))}: "
                f"{_fmt_currency(pm.get('amount'), proposal.currency)}"
            )

    if content.get("terms_and_conditions"):
        document.add_heading("Terms & Conditions", level=2)
        document.add_paragraph(_safe_text(content["terms_and_conditions"]))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()
